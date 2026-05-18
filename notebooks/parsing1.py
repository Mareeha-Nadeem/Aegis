from pathlib import Path
import json
import asyncio
from concurrent.futures import ThreadPoolExecutor
import re

import fitz
import markdown
from bs4 import BeautifulSoup
from docx import Document
import mammoth

# =====================================
# PATHS
# =====================================

BASE_DIR = Path(__file__).resolve().parents[1]

DOCUMENTS_DIR = BASE_DIR / "data" / "raw"

# JSON OUTPUT (instead of Excel)
OUTPUT_FILE = BASE_DIR / "data" / "documents" / "all_parsed.json"

TRACKER = BASE_DIR / "data" / "documents" / "processed.json"
TRACKER.parent.mkdir(parents=True, exist_ok=True)

OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)


# =====================================
# CLEAN TEXT (RAG SAFE VERSION)
# =====================================

def clean_text(text: str) -> str:
    if not isinstance(text, str):
        return ""

    # remove control chars (PDF garbage)
    text = re.sub(r"[\x00-\x1F\x7F]", " ", text)

    # remove weird symbols but keep readable sentence structure
    text = re.sub(r"[^\w\s.,;:()\-/'\"%&@!?]", " ", text)

    # normalize spaces
    text = re.sub(r"\s+", " ", text)

    # optional: remove super short junk lines
    return text.strip()


# =====================================
# LOAD TRACKER
# =====================================

if TRACKER.exists():
    try:
        processed = set(json.loads(TRACKER.read_text(encoding="utf-8")))
    except:
        processed = set()
else:
    processed = set()

SUPPORTED_EXTENSIONS = {".pdf", ".md", ".markdown", ".html", ".htm", ".txt",".docx"}


# =====================================
# PARSERS
# =====================================

def parse_pdf(path):
    doc = fitz.open(path)
    text = []

    for page in doc:
        page_text = page.get_text("text")
        if page_text:
            text.append(page_text)

    return clean_text(" ".join(text))


def parse_markdown(path):
    content = path.read_text(encoding="utf-8", errors="ignore")
    html = markdown.markdown(content)
    soup = BeautifulSoup(html, "html.parser")
    return clean_text(soup.get_text())


def parse_html(path):
    content = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(content, "html.parser")
    return clean_text(soup.get_text())


def parse_txt(path):
    content = path.read_text(encoding="utf-8", errors="ignore")
    return clean_text(content)

# def parse_docx(path):
#     doc = Document(path)

#     text = []
#     for para in doc.paragraphs:
#         text.append(para.text)

#     return clean_text(" ".join(text))


def parse_docx(path):
    try:
        doc = Document(path)

        text = []
        for para in doc.paragraphs:
            text.append(para.text)

        return clean_text(" ".join(text))

    except Exception:
        import textract

        text = textract.process(str(path))
        return clean_text(text.decode("utf-8", errors="ignore"))


def parse_doc(path):
    import textract

    text = textract.process(str(path))

    return clean_text(text.decode("utf-8", errors="ignore"))

# def parse_file(path):
#     ext = path.suffix.lower()

#     if ext == ".pdf":
#         return parse_pdf(path)
#     elif ext in [".md", ".markdown"]:
#         return parse_markdown(path)
#     elif ext in [".html", ".htm"]:
#         return parse_html(path)
#     elif ext == ".txt":
#         return parse_txt(path)

#     return None

def parse_file(path):
    ext = path.suffix.lower()

    if ext == ".pdf":
        return parse_pdf(path)

    elif ext in [".md", ".markdown"]:
        return parse_markdown(path)

    elif ext in [".html", ".htm"]:
        return parse_html(path)

    elif ext == ".txt":
        return parse_txt(path)

    elif ext == ".docx":
        return parse_docx(path)

    elif ext == ".doc":
        return parse_doc(path)

    return None


# =====================================
# THREADING
# =====================================

executor = ThreadPoolExecutor(max_workers=6)


async def process_file(loop, file):
    file_id = str(file)

    if file_id in processed:
        return None

    try:
        text = await loop.run_in_executor(executor, parse_file, file)

        if text:
            processed.add(file_id)

            return {
                "file_name": file.name,
                "file_path": file_id,
                "content": text
            }

    except Exception as e:
        print("ERROR:", file.name, e)

    return None


# =====================================
# MAIN
# =====================================

# async def main():
#     loop = asyncio.get_running_loop()

#     files = [
#         f for f in DOCUMENTS_DIR.rglob("*")
#         if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
#     ]

#     tasks = [process_file(loop, f) for f in files]
#     results = await asyncio.gather(*tasks)

#     new_rows = [r for r in results if r]

async def main():
    loop = asyncio.get_running_loop()

    files = [
        f for f in DOCUMENTS_DIR.rglob("*")
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ]

    tasks = [process_file(loop, f) for f in files]
    results = await asyncio.gather(*tasks)

    new_rows = [r for r in results if r]

    # LOAD OLD JSON
    if OUTPUT_FILE.exists():
        try:
            existing_data = json.loads(
                OUTPUT_FILE.read_text(encoding="utf-8")
            )
        except:
            existing_data = []
    else:
        existing_data = []

    # ADD IDS
    start_id = len(existing_data) + 1

    for i, row in enumerate(new_rows, start=start_id):
        row["id"] = i
    # =====================================
    # LOAD OLD JSON
    # =====================================

    # if OUTPUT_FILE.exists():
    #     try:
    #         existing_data = json.loads(OUTPUT_FILE.read_text(encoding="utf-8"))
    #     except:
    #         existing_data = []
    # else:
    #     existing_data = []


    # =====================================
    # MERGE DATA
    # =====================================

    if new_rows:
        final_data = existing_data + new_rows

        # save tracker first (safe point)
        TRACKER.write_text(
            json.dumps(list(processed), indent=2),
            encoding="utf-8"
        )

        # save clean JSON
        OUTPUT_FILE.write_text(
            json.dumps(final_data, indent=2, ensure_ascii=False),
            encoding="utf-8"
        )

        print("SAVED JSON:", OUTPUT_FILE)

    else:
        print("NO NEW FILES")

        TRACKER.write_text(
            json.dumps(list(processed), indent=2),
            encoding="utf-8"
        )


# =====================================
# RUN
# =====================================

if __name__ == "__main__":
    asyncio.run(main())